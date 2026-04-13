import os
import json
from datetime import datetime
from openai import OpenAI
from docx import Document
from dotenv import load_dotenv

load_dotenv()

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

def validate_input(brief_text):
    """Check if brief has minimum required content"""
    if not brief_text or len(brief_text.strip()) < 50:
        return False, "Brief too short (minimum 50 characters)"
    return True, "Valid"

def call_ai_to_generate_document(brief_text):
    """Send brief to OpenAI and get structured document content"""
    
    system_prompt = """You are a civil engineering project planner. 
    Convert project briefs into structured documents with these EXACT sections:
    
    1. EXECUTIVE SUMMARY (2-3 sentences)
    2. PROJECT SCOPE (bullet points)
    3. KEY DELIVERABLES (numbered list)
    4. RISKS & CONSTRAINTS (bullet points)
    5. NEXT ACTIONS (numbered list)
    
    Return ONLY valid JSON with this structure:
    {
        "executive_summary": "...",
        "project_scope": ["item1", "item2"],
        "key_deliverables": ["1. item", "2. item"],
        "risks": ["risk1", "risk2"],
        "next_actions": ["1. action", "2. action"]
    }"""
    
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": brief_text}
            ],
            temperature=0.3,
            response_format={"type": "json_object"}
        )
        
        content = response.choices[0].message.content
        return json.loads(content) # type: ignore
        
    except Exception as e:
        print(f"AI API Error: {e}")
        return None

def create_word_document(data, project_name):
    """Generate Word document from structured data"""
    
    doc = Document()
    
    # Title
    doc.add_heading(f"Project Plan: {project_name}", 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(data.get("executive_summary", "Not specified"))
    
    # Project Scope
    doc.add_heading("2. Project Scope", level=1)
    for item in data.get("project_scope", []):
        doc.add_paragraph(f"• {item}", style='List Bullet')
    
    # Key Deliverables
    doc.add_heading("3. Key Deliverables", level=1)
    for item in data.get("key_deliverables", []):
        doc.add_paragraph(item, style='List Number')
    
    # Risks & Constraints
    doc.add_heading("4. Risks & Constraints", level=1)
    for risk in data.get("risks", []):
        doc.add_paragraph(f"• {risk}", style='List Bullet')
    
    # Next Actions
    doc.add_heading("5. Next Actions", level=1)
    for action in data.get("next_actions", []):
        doc.add_paragraph(action, style='List Number')
    
    # Save
    filename = f"{project_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    return filename

def main():
    print("=" * 50)
    print("TA Project Services - AI Document Generator")
    print("=" * 50)
    
    # Input
    project_name = input("Project name: ").strip()
    print("\nEnter project brief (press Enter twice when done):")
    lines = []
    while True:
        line = input()
        if line == "" and lines and lines[-1] == "":
            break
        lines.append(line)
    
    brief_text = "\n".join(lines)
    
    # Validation
    is_valid, message = validate_input(brief_text)
    if not is_valid:
        print(f"\n❌ Error: {message}")
        return
    
    print(f"\n✅ Brief validated ({len(brief_text)} characters)")
    print("🤖 Generating document with AI...")
    
    # AI processing
    ai_data = call_ai_to_generate_document(brief_text)
    if not ai_data:
        print("❌ AI processing failed")
        return
    
    # Create document
    filename = create_word_document(ai_data, project_name)
    print(f"✅ Word document created: {filename}")
    
    # Show preview
    print("\n" + "=" * 50)
    print("DOCUMENT PREVIEW:")
    print("=" * 50)
    print(f"Executive Summary:\n{ai_data['executive_summary'][:200]}...")
    print(f"\nScope items: {len(ai_data.get('project_scope', []))}")
    print(f"Risks identified: {len(ai_data.get('risks', []))}")
    print("=" * 50)
    
    # Save for next steps
    with open("last_run.json", "w") as f:
        json.dump({
            "project_name": project_name,
            "filename": filename,
            "ai_data": ai_data,
            "timestamp": datetime.now().isoformat()
        }, f)
    
    print("\n✅ Ready for Asana + OneDrive integration")

if __name__ == "__main__":
    main()